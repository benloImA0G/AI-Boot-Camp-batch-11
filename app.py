# ============================================================
# ResumeFit AI v5.0 — Hybrid: TF-IDF + Pinecone + LLM
# Author: Bernard Lokasasmita
# Bootcamp: Ruangguru AI Engineering Batch 11
# ============================================================

import streamlit as st
import pypdf
import pandas as pd
import numpy as np
import re
import io
import os
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import plotly.graph_objects as go
import plotly.express as px
from openai import OpenAI
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── NEW: Vector DB imports ──────────────────────────────────
from sentence_transformers import SentenceTransformer
from pinecone import Pinecone, ServerlessSpec

# PAGE CONFIG
st.set_page_config(
    page_title="ResumeFit AI v5.0",
    page_icon="🚀",
    layout="wide",
    initial_sidebar_state="expanded"
)

# CUSTOM CSS
st.markdown("""
<style>
    .main { background-color: #f8fafc; }
    .stButton>button {
        background: linear-gradient(135deg, #4f0d8f, #7c3aed);
        color: white; border: none; border-radius: 8px;
        padding: 12px 28px; font-size: 15px; font-weight: 700;
        width: 100%; cursor: pointer;
    }
    .stButton>button:hover { opacity: 0.9; }
    .header-box {
        background: linear-gradient(135deg, #4f0d8f, #7c3aed);
        padding: 28px 32px; border-radius: 12px; margin-bottom: 24px;
    }
    .aspect-card {
        background: white; border-radius: 10px; padding: 16px 20px;
        box-shadow: 0 2px 8px rgba(0,0,0,0.07); margin-bottom: 12px;
        border-left: 5px solid #7c3aed;
    }
    .aspect-card h4 { margin: 0 0 6px 0; color: #4f0d8f; font-size: 1rem; }
    .aspect-card p  { margin: 0; color: #374151; font-size: 0.88rem; }
    .score-badge-green  { background:#dcfce7; color:#166534; border-radius:20px; padding:3px 12px; font-size:0.82rem; font-weight:700; }
    .score-badge-yellow { background:#fef9c3; color:#854d0e; border-radius:20px; padding:3px 12px; font-size:0.82rem; font-weight:700; }
    .score-badge-red    { background:#fee2e2; color:#991b1b; border-radius:20px; padding:3px 12px; font-size:0.82rem; font-weight:700; }
    .skill-found   { display:inline-block; background:#dcfce7; color:#166534; border-radius:20px; padding:4px 12px; margin:3px; font-size:0.83rem; font-weight:600; }
    .skill-missing { display:inline-block; background:#fee2e2; color:#991b1b; border-radius:20px; padding:4px 12px; margin:3px; font-size:0.83rem; font-weight:600; }
    .before-box { background:#fff7ed; border-left:4px solid #f97316; padding:12px 16px; border-radius:6px; font-size:0.88rem; color:#374151; white-space:pre-wrap; }
    .after-box  { background:#f0fdf4; border-left:4px solid #22c55e; padding:12px 16px; border-radius:6px; font-size:0.88rem; color:#374151; white-space:pre-wrap; }
    .section-header {
        background: linear-gradient(135deg, #4f0d8f, #7c3aed);
        color: white; padding: 10px 18px; border-radius: 8px;
        font-size: 1.05rem; font-weight: 700; margin: 18px 0 10px 0;
    }
    .metric-card {
        background: white; border-radius: 12px; padding: 18px;
        box-shadow: 0 2px 8px rgba(0,0,0,0.08); text-align: center;
        border-left: 5px solid #7c3aed;
    }
    .metric-card h1 { font-size: 2.2rem; color: #4f0d8f; margin: 0; }
    .metric-card p  { color: #6b7280; margin: 4px 0 0 0; font-size: 0.85rem; }
    .error-box   { background:#fee2e2; color:#991b1b; padding:14px 18px; border-radius:8px; font-size:0.88rem; margin-top:8px; border-left:4px solid #ef4444; }
    .success-box { background:#dcfce7; color:#166534; padding:14px 18px; border-radius:8px; font-size:0.88rem; margin-top:8px; border-left:4px solid #22c55e; }
    .info-box    { background:#eff6ff; color:#1e40af; padding:14px 18px; border-radius:8px; font-size:0.88rem; margin-top:8px; border-left:4px solid #3b82f6; }
    .warn-box    { background:#fef9c3; color:#854d0e; padding:14px 18px; border-radius:8px; font-size:0.88rem; margin-top:8px; border-left:4px solid #f59e0b; }
    .rule-box    { background:#1e1e2e; color:#a6e3a1; padding:14px 18px; border-radius:8px; font-family:monospace; font-size:0.82rem; white-space:pre-wrap; margin-top:8px; }
    .strategy-box { background:#f5f3ff; border-left:4px solid #7c3aed; padding:14px 18px; border-radius:8px; font-size:0.9rem; color:#4f0d8f; margin:12px 0; }
    .vector-box  { background:#f0f9ff; border-left:4px solid #0ea5e9; padding:14px 18px; border-radius:8px; font-size:0.88rem; color:#0c4a6e; margin-top:8px; }
</style>
""", unsafe_allow_html=True)

# ── CONSTANTS ──────────────────────────────────────────────
SKILL_DICT = [
    # ── Programming Languages ──────────────────────────────
    "python","java","javascript","typescript","c++","c#","r","scala","go","rust",
    "kotlin","swift","php","ruby","matlab","julia","dart","elixir","haskell",

    # ── Web & App Frameworks ───────────────────────────────
    "react","nextjs","vue","angular","svelte","tailwind","html","css",
    "fastapi","flask","django","express","spring boot","laravel","rails",

    # ── Classic ML / Data Science ──────────────────────────
    "machine learning","deep learning","nlp","natural language processing",
    "computer vision","statistical modeling","time series","forecasting",
    "feature engineering","data wrangling","eda","exploratory data analysis",
    "tensorflow","pytorch","keras","scikit-learn","xgboost","lightgbm","catboost",
    "pandas","numpy","scipy","statsmodels",

    # ── Generative AI & LLM (expanded) ────────────────────
    "generative ai","large language model","llm","gpt","chatgpt","claude",
    "gemini","llama","mistral","openai","anthropic","cohere","groq",
    "hugging face","transformers","langchain","llamaindex","langgraph",
    "prompt engineering","few-shot","zero-shot","chain of thought",
    "rag","retrieval augmented generation","agentic ai","ai agent",
    "fine-tuning","lora","qlora","peft","instruction tuning",
    "multimodal","text to image","stable diffusion","midjourney",
    "whisper","speech to text","text to speech",

    # ── Vector DB & Embeddings ─────────────────────────────
    "vector database","vector search","semantic search","embeddings",
    "pinecone","faiss","weaviate","qdrant","chroma","milvus","pgvector",

    # ── MLOps & Model Deployment ──────────────────────────
    "mlops","model deployment","model serving","model monitoring",
    "mlflow","wandb","weights and biases","kubeflow","bentoml",
    "sagemaker","vertex ai","azure ml","hugging face hub",
    "api","rest api","grpc","model registry","a/b testing",

    # ── Databases ─────────────────────────────────────────
    "sql","mysql","postgresql","mongodb","redis","elasticsearch",
    "snowflake","bigquery","supabase","firebase","dynamodb","cassandra",

    # ── Data Engineering ──────────────────────────────────
    "spark","hadoop","airflow","dbt","kafka","flink","databricks",
    "etl","data pipeline","data warehouse","data lake","data mesh",
    "looker","metabase","data catalog","data governance",

    # ── Cloud & DevOps ─────────────────────────────────────
    "aws","gcp","azure","docker","kubernetes","ci/cd","terraform",
    "linux","bash","git","github","gitlab","vercel","netlify",
    "microservices","serverless","cloud native","devops","sre",

    # ── Visualization & BI ─────────────────────────────────
    "tableau","power bi","streamlit","plotly","matplotlib","seaborn",
    "looker studio","grafana","d3.js",

    # ── Automation & No-Code ──────────────────────────────
    "n8n","zapier","make","automation","rpa","power automate",
    "opencv","yolo","object detection",

    # ── Soft Skills ───────────────────────────────────────
    "communication","leadership","teamwork","collaboration",
    "problem solving","critical thinking","analytical thinking",
    "time management","presentation","public speaking","mentoring",
    "project management","agile","scrum","kanban","okr",
    "stakeholder management","cross functional","negotiation",

    # ── Business & Product ────────────────────────────────
    "product management","product roadmap","go-to-market","brand strategy",
    "market research","competitor analysis","sales strategy",
    "product launch","pricing strategy","user research","ux","ui design",
    "figma","wireframing","prototyping","growth hacking","seo","sem",

    # ── Office & Productivity ─────────────────────────────
    "excel","powerpoint","word","google sheets","google slides",
    "jira","confluence","notion","trello","asana","slack",
    "microsoft office","google workspace",

    # ── Indonesian Job Market Specific ────────────────────
    "bahasa indonesia","business development","account management",
    "digital marketing","social media","content creation","copywriting",
    "e-commerce","tokopedia","shopee","gojek","grab",
    "kpi","reporting","budgeting","financial modeling",
]

ASPECTS = [
    ("Overall Impression",          "🌟", "General quality and first impression of the CV"),
    ("Contact Information",         "📞", "Completeness of contact details"),
    ("Relevant Skills",             "🛠️",  "Technical and soft skills alignment"),
    ("Professional Summary",        "📝", "Quality of the professional summary/objective"),
    ("Work Experience",             "💼", "Relevance and impact of work experience"),
    ("Achievements",                "🏆", "Quantified accomplishments and impact"),
    ("Education & Certification",   "🎓", "Educational background and certifications"),
    ("Organizational Activity",     "🤝", "Extracurricular and organizational involvement"),
    ("Consistent & Error-free",     "✅", "Grammar, formatting, and consistency"),
    ("Additional Sections",         "➕", "Portfolio, projects, languages, hobbies"),
    ("Keywords",                    "🔑", "ATS keyword optimization"),
    ("Career Recommendation",       "🎯", "Suggested career paths and roles"),
]

GROQ_BASE_URL      = "https://api.groq.com/openai/v1"
GROQ_DEFAULT_MODEL = "llama-3.1-8b-instant"
OPENAI_MODELS      = ["gpt-3.5-turbo", "gpt-4o-mini", "gpt-4o"]

# ── Feature #1: Impact Language constants ───────────────────
WEAK_VERBS = [
    "responsible for", "helped", "assisted", "worked on", "involved in",
    "participated in", "supported", "contributed to", "was part of",
    "tasked with", "duties included", "in charge of", "handled",
]
BUZZWORDS = [
    "results-driven", "team player", "hard worker", "synergy", "go-getter",
    "passionate", "detail-oriented", "self-starter", "dynamic", "proactive",
    "motivated", "enthusiastic", "innovative thinker", "thought leader",
    "rockstar", "ninja", "guru", "visionary", "disruptive",
]

# ── Feature #2: ATS formatting constants ────────────────────
REQUIRED_SECTIONS = ["experience", "education", "skills"]
ATS_SCORE_WEIGHTS = {"images": 25, "columns": 20, "tables": 15,
                     "sections": 20, "contact": 10, "length": 10}

# ── NEW: Vector DB constants ────────────────────────────────
EMBEDDING_MODEL_NAME = "all-MiniLM-L6-v2"   # 384 dimensions, free, no API key
PINECONE_INDEX_NAME  = "resumefit-index"
PINECONE_DIMENSION   = 384
PINECONE_METRIC      = "cosine"

# ── LLM CLIENTS ────────────────────────────────────────────
def get_groq_client(api_key):
    return OpenAI(api_key=api_key, base_url=GROQ_BASE_URL)

def get_openai_client(api_key):
    return OpenAI(api_key=api_key)

# ── NEW: Embedding model (cached — loads once per session) ──
@st.cache_resource
def load_embedding_model():
    return SentenceTransformer(EMBEDDING_MODEL_NAME)

# ── NEW: Pinecone init ──────────────────────────────────────
def init_pinecone(api_key: str):
    """
    Connect to Pinecone and return the index.
    Creates the index automatically if it doesn't exist yet.
    """
    pc = Pinecone(api_key=api_key)
    existing = [idx.name for idx in pc.list_indexes()]
    if PINECONE_INDEX_NAME not in existing:
        pc.create_index(
            name=PINECONE_INDEX_NAME,
            dimension=PINECONE_DIMENSION,
            metric=PINECONE_METRIC,
            spec=ServerlessSpec(cloud="aws", region="us-east-1")
        )
    return pc.Index(PINECONE_INDEX_NAME)

# ── NEW: Step 1 — Program Ingestion ────────────────────────
def ingest_job_description(jd_text: str, pinecone_index, embed_model) -> list:
    """
    Program Ingestion (matches mentor's diagram):
    Job Description → Model Embedding → upsert into Vector Database.
    Returns the JD embedding vector.
    """
    jd_vector = embed_model.encode(jd_text[:2000]).tolist()
    pinecone_index.upsert(vectors=[{
        "id":     "job_desc_current",
        "values": jd_vector
    }])
    return jd_vector

# ── NEW: Step 2 — Model Embedding + Query ──────────────────
def query_with_cv(cv_text: str, pinecone_index, embed_model) -> float:
    """
    Model Embedding + Query (matches mentor's diagram):
    CV Text → Model Embedding → query Vector Database → similarity score (0–100).
    """
    try:
        cv_vector = embed_model.encode(cv_text[:2000]).tolist()
        result = pinecone_index.query(
            vector=cv_vector,
            top_k=1,
            include_values=False
        )
        if result.matches:
            return round(result.matches[0].score * 100, 1)
        return 0.0
    except Exception as e:
        st.markdown(
            f"<div class='warn-box'>⚠️ Pinecone query failed: {e}<br>"
            "Falling back to TF-IDF only.</div>",
            unsafe_allow_html=True
        )
        return 0.0

# ── Feature #1: Impact Language Analysis ───────────────────
def analyze_impact_language(resume_text: str, client=None, model: str = ""):
    """
    Two-layer analysis:
    - Rule-based: flag weak verbs & buzzwords instantly
    - LLM: deeper impact score, top tip, quantification check
    """
    rt_lower = resume_text.lower()

    # Rule-based layer
    found_weak = [v for v in WEAK_VERBS if v in rt_lower]
    found_buzz = [b for b in BUZZWORDS if b in rt_lower]
    has_numbers = bool(re.search(
        r'\d+\s*%|\$\s*\d+|\d+\s*x\b|\d+\s*million|\d+\s*billion|'
        r'\d+\s*team|\d+\s*users|\d+\s*clients|\d+\s*years',
        rt_lower
    ))

    llm_result = {"score": None, "tip": None, "quantified": None}

    if client:
        prompt = (
            f"Analyze this CV for impact language quality.\n\n"
            f"RESUME:\n{resume_text[:1500]}\n\n"
            "Output EXACTLY these 4 lines — plain text only, no markdown:\n"
            "IMPACT_SCORE: <0-100>\n"
            "QUANTIFIED: <Yes or No>\n"
            "WEAK_PHRASE: <single most damaging weak phrase found, or None>\n"
            "TOP_TIP: <one specific actionable sentence to improve impact>\n"
        )
        try:
            r = client.chat.completions.create(
                model=model,
                messages=[{"role": "user", "content": prompt}],
                max_tokens=150,
                temperature=0.1
            )
            raw = r.choices[0].message.content
            for line in raw.split("\n"):
                l = line.strip()
                if l.lower().startswith("impact_score:"):
                    try: llm_result["score"] = int(re.sub(r'[^0-9]', '', l.split(":", 1)[1]))
                    except: pass
                elif l.lower().startswith("quantified:"):
                    llm_result["quantified"] = "yes" in l.lower()
                elif l.lower().startswith("top_tip:"):
                    llm_result["tip"] = l.split(":", 1)[1].strip()
        except:
            pass

    return found_weak, found_buzz, has_numbers, llm_result


# ── Feature #2: ATS Formatting Checks ──────────────────────
def check_ats_formatting(uploaded_file) -> list:
    """
    PDF structure checks via pypdf:
    images, encryption, page count, embedded fonts.
    Returns list of (issue_label, detail, severity) tuples.
    severity: 'error' | 'warning' | 'ok'
    """
    issues = []
    try:
        uploaded_file.seek(0)
        reader = pypdf.PdfReader(uploaded_file)

        # Encryption
        if reader.is_encrypted:
            issues.append(("🔒 Encrypted PDF", "Password-protected — ATS cannot read it at all.", "error"))

        # Images / graphics
        total_images = sum(len(page.images) for page in reader.pages)
        if total_images > 0:
            issues.append(("🖼️ Images detected",
                           f"{total_images} image(s) found — ATS skips images entirely. "
                           "Remove photos, logos, and icon-based skill bars.", "error"))
        else:
            issues.append(("🖼️ No images", "Clean — no images detected.", "ok"))

        # Page count
        pages = len(reader.pages)
        if pages > 2:
            issues.append(("📄 Too long",
                           f"{pages} pages — most ATS and recruiters prefer 1–2 pages.", "warning"))
        else:
            issues.append(("📄 Page count", f"{pages} page(s) — good length.", "ok"))

        uploaded_file.seek(0)
    except Exception as e:
        issues.append(("⚠️ PDF read error", str(e), "error"))
    return issues


def check_text_ats(resume_text: str) -> list:
    """
    Text-pattern checks for ATS-breaking formatting:
    columns, tables, special chars, missing sections, contact info.
    """
    issues = []
    lines = [l for l in resume_text.split("\n") if l.strip()]

    # Multi-column detection: many short lines side-by-side
    short_lines = [l for l in lines if 0 < len(l.strip()) < 35]
    col_ratio = len(short_lines) / max(len(lines), 1)
    if col_ratio > 0.45:
        issues.append(("⚠️ Possible multi-column layout",
                       f"{int(col_ratio*100)}% of lines are very short — "
                       "multi-column CVs often fail ATS parsing. Use single-column.", "error"))
    else:
        issues.append(("✅ Single-column layout", "No multi-column pattern detected.", "ok"))

    # Table detection: pipe characters
    pipe_count = resume_text.count("|")
    if pipe_count > 2:
        issues.append(("⚠️ Table detected",
                       f"{pipe_count} pipe characters found — tables break ATS text extraction. "
                       "Replace with plain bullet lists.", "error"))
    else:
        issues.append(("✅ No tables", "No table patterns detected.", "ok"))

    # Special / non-ASCII characters
    special = len(re.findall(r'[^\x00-\x7F]', resume_text))
    if special > 15:
        issues.append(("⚠️ Special characters",
                       f"{special} non-standard characters — may cause ATS parsing errors. "
                       "Replace with standard ASCII equivalents.", "warning"))
    else:
        issues.append(("✅ Clean characters", "No problematic special characters.", "ok"))

    # Required sections
    missing_secs = [s for s in REQUIRED_SECTIONS if s not in resume_text.lower()]
    if missing_secs:
        issues.append(("⚠️ Missing sections",
                       f"Cannot detect: {', '.join(missing_secs)}. "
                       "ATS looks for these headings explicitly.", "warning"))
    else:
        issues.append(("✅ Key sections present",
                       "Experience, Education, and Skills sections detected.", "ok"))

    # Contact info
    has_email = "@" in resume_text
    has_phone = bool(re.search(r'\+?\d[\d\s\-().]{7,}', resume_text))
    if not has_email:
        issues.append(("⚠️ No email address", "Email not detected — always include it.", "error"))
    else:
        issues.append(("✅ Email present", "Email address detected.", "ok"))
    if not has_phone:
        issues.append(("⚠️ No phone number", "Phone number not detected.", "warning"))
    else:
        issues.append(("✅ Phone present", "Phone number detected.", "ok"))

    return issues


def ats_format_score(pdf_issues: list, text_issues: list) -> int:
    """Compute an ATS Format Score (0–100) from issue severity counts."""
    errors   = sum(1 for _, _, s in pdf_issues + text_issues if s == "error")
    warnings = sum(1 for _, _, s in pdf_issues + text_issues if s == "warning")
    score = max(0, 100 - (errors * 20) - (warnings * 8))
    return score


# ── HELPERS ────────────────────────────────────────────────
def extract_text_from_pdf(file):
    try:
        reader = pypdf.PdfReader(file)
        text = ""
        for page in reader.pages:
            t = page.extract_text()
            if t:
                text += t + "\n"
        return text.strip()
    except Exception as e:
        st.error(f"PDF error: {e}")
        return ""

def clean_text(text):
    text = text.lower()
    text = re.sub(r'[^a-z0-9\s+#]', ' ', text)
    return re.sub(r'\s+', ' ', text).strip()

def extract_skills(text):
    tl = text.lower()
    return list({s for s in SKILL_DICT if s.lower() in tl})

def calculate_tfidf(resume, job):
    vec = TfidfVectorizer(stop_words='english', ngram_range=(1,2))
    try:
        m = vec.fit_transform([clean_text(resume), clean_text(job)])
        return round(cosine_similarity(m[0:1], m[1:2])[0][0] * 100, 1)
    except:
        return 0.0

def keyword_coverage(resume, job):
    jw = {w for w in clean_text(job).split() if len(w) > 3}
    rw = set(clean_text(resume).split())
    if not jw:
        return 0.0, [], []
    matched = jw & rw
    missing = jw - rw
    return round(len(matched)/len(jw)*100, 1), list(matched)[:15], list(missing)[:15]

def skill_analysis(resume, job):
    rs = extract_skills(resume)
    js = extract_skills(job)
    return rs, js, list(set(rs)&set(js)), list(set(js)-set(rs)), list(set(rs)-set(js))

def score_label(s):
    if s >= 75: return "Excellent Match", "#22c55e"
    if s >= 55: return "Good Match",      "#3b82f6"
    if s >= 35: return "Moderate Match",  "#f59e0b"
    return "Needs Improvement", "#ef4444"

def badge(score):
    if score >= 75: return "<span class='score-badge-green'>Strong ✅</span>"
    if score >= 50: return "<span class='score-badge-yellow'>Average ⚠️</span>"
    return "<span class='score-badge-red'>Weak ❌</span>"

def clean_ai_text(text):
    text = re.sub(r"\*{1,3}(.*?)\*{1,3}", r"\1", text)
    text = re.sub(r"`{1,3}(.*?)`{1,3}", r"\1", text)
    text = re.sub(r"^#{1,6}\s*", "", text, flags=re.MULTILINE)
    return text.strip()

# ── CHARTS ─────────────────────────────────────────────────
def gauge_chart(score):
    label, color = score_label(score)
    fig = go.Figure(go.Indicator(
        mode="gauge+number",
        value=score,
        title={"text": f"ATS Match Score<br><span style='font-size:0.8em;color:{color}'>{label}</span>", "font":{"size":15}},
        gauge={
            "axis": {"range":[0,100]},
            "bar":  {"color": color},
            "steps":[
                {"range":[0,35],  "color":"#fee2e2"},
                {"range":[35,55], "color":"#fef9c3"},
                {"range":[55,75], "color":"#dbeafe"},
                {"range":[75,100],"color":"#dcfce7"},
            ],
        }
    ))
    fig.update_layout(height=260, margin=dict(t=60,b=10,l=10,r=10))
    return fig

def radar_chart(matched, missing):
    all_s = (matched+missing)[:8]
    if len(all_s) < 3: return None
    rv = [1 if s in matched else 0 for s in all_s]
    fig = go.Figure()
    fig.add_trace(go.Scatterpolar(r=[1]*len(all_s)+[1], theta=all_s+[all_s[0]], fill="toself", name="Job Req", fillcolor="rgba(124,58,237,0.15)", line=dict(color="#7c3aed",width=2)))
    fig.add_trace(go.Scatterpolar(r=rv+[rv[0]], theta=all_s+[all_s[0]], fill="toself", name="Your CV", fillcolor="rgba(34,197,94,0.2)", line=dict(color="#22c55e",width=2)))
    fig.update_layout(polar=dict(radialaxis=dict(visible=True,range=[0,1])), showlegend=True, height=300, margin=dict(t=40,b=30,l=30,r=30), title="Skill Radar")
    return fig

def aspect_bar_chart(scores_dict):
    aspects = list(scores_dict.keys())
    scores  = [v if isinstance(v, (int,float)) else v.get("score",0) for v in scores_dict.values()]
    colors  = ["#22c55e" if s>=75 else "#f59e0b" if s>=50 else "#ef4444" for s in scores]
    fig = go.Figure(go.Bar(x=scores, y=aspects, orientation='h', marker_color=colors, text=[f"{s}%" for s in scores], textposition="auto"))
    fig.update_layout(height=420, margin=dict(t=20,b=20,l=10,r=10), xaxis=dict(range=[0,100]), title="12-Aspect CV Score Breakdown")
    return fig

# ── NEW: Scoring comparison chart ──────────────────────────
def scoring_breakdown_chart(tfidf_score, vector_score, cov_pct, skill_score):
    methods = ["TF-IDF<br>(Statistical)", "Pinecone<br>(Semantic)", "Keyword<br>(Coverage)", "Skill<br>(Match)"]
    scores  = [tfidf_score, vector_score, cov_pct, skill_score]
    weights = ["40%", "30%", "20%", "10%"]
    colors  = ["#7c3aed", "#0ea5e9", "#f59e0b", "#22c55e"]
    fig = go.Figure(go.Bar(
        x=methods, y=scores,
        marker_color=colors,
        text=[f"{s}%<br><small>({w})</small>" for s, w in zip(scores, weights)],
        textposition="outside"
    ))
    fig.update_layout(
        title="4-Layer Hybrid Scoring Breakdown",
        yaxis=dict(range=[0,110]),
        height=320,
        margin=dict(t=50,b=20,l=10,r=10)
    )
    return fig

# ── AI FUNCTIONS ───────────────────────────────────────────
def test_api_connection(client, model):
    try:
        r = client.chat.completions.create(
            model=model,
            messages=[{"role": "user", "content": "Say OK"}],
            max_tokens=5
        )
        _ = r.choices[0].message.content
        return True, f"Connected ✅  (model: {model})"
    except Exception as e:
        return False, str(e)

def analyze_12_aspects(resume_text, job_text, client, model):
    system_msg = (
        "You are a Strategic Career Architect. Provide a brutally honest, specific gap analysis "
        "between this candidate and the job. Be concrete — reference actual content from the resume. "
        "No markdown, no bold, no bullet points. Plain text only."
    )
    user_msg = (
        f"RESUME:\n{resume_text[:2000]}\n\n"
        f"JOB DESCRIPTION:\n{job_text[:1200]}\n\n"
        "Analyze ALL 12 aspects. For each, output EXACTLY this block:\n"
        "ASPECT: <name>\n"
        "SCORE: <0-100>\n"
        "ASSESSMENT: <one specific sentence referencing the resume>\n"
        "TIP: <one concrete actionable improvement>\n"
        "---\n\n"
        "ASPECTS (all 12 in order):\n"
        "1. Overall Impression\n2. Contact Information\n3. Relevant Skills\n"
        "4. Professional Summary\n5. Work Experience\n6. Achievements\n"
        "7. Education & Certification\n8. Organizational Activity\n"
        "9. Consistent & Error-free Writing\n10. Additional Sections\n"
        "11. Keywords\n12. Career Recommendation\n\n"
        "RULES: Plain text only. No **, no *, no #. All 12 blocks required. No intro or closing text."
    )
    try:
        r = client.chat.completions.create(
            model=model,
            messages=[{"role": "system", "content": system_msg}, {"role": "user", "content": user_msg}],
            max_tokens=2800,
            temperature=0.2
        )
        return r.choices[0].message.content, None
    except Exception as e:
        return "", str(e)

def parse_aspect_scores(raw_text):
    results = {}
    if not raw_text:
        return results
    blocks = re.split(r"\n\s*-{2,}\s*\n|\n\s*-{2,}\s*$|^\s*-{2,}\s*$", raw_text, flags=re.MULTILINE)
    for block in blocks:
        if not block.strip():
            continue
        name = score = assessment = tip = None
        for line in block.split("\n"):
            l = line.strip().replace("**","").replace("`","")
            m = re.match(r"(?i)^aspect\s*[:\-]\s*(.+)$", l)
            if m: name = m.group(1).strip()
            m = re.match(r"(?i)^score\s*[:\-]\s*([0-9]{1,3})", l)
            if m: score = m.group(1).strip()
            m = re.match(r"(?i)^assessment\s*[:\-]\s*(.+)$", l)
            if m: assessment = m.group(1).strip()
            m = re.match(r"(?i)^tip\s*[:\-]\s*(.+)$", l)
            if m: tip = m.group(1).strip()
        if name and score:
            try:
                s = max(0, min(100, int(re.sub(r'[^0-9]', '', score))))
                results[name] = {"score": s, "assessment": assessment or "", "tip": tip or ""}
            except:
                pass
    return results

def rewrite_cv_sections(resume_text, job_text, missing_skills, final_score, client, model):
    ms = ', '.join(missing_skills[:10]) if missing_skills else 'None'

    if final_score >= 60:
        strategy = "HIGH_MATCH"
        strategy_instruction = (
            f"The candidate has a {final_score}% match with this job. "
            "STRATEGY: Tailor the CV to this specific job. "
            "Align wording with the job description. Emphasize matching skills. "
            "Incorporate missing keywords naturally where truthful."
        )
    else:
        strategy = "LOW_MATCH"
        strategy_instruction = (
            f"The candidate has only a {final_score}% match with this job — the CV is NOT a strong fit. "
            "STRATEGY: Do NOT force-fit the candidate into this role. "
            "Instead, ENHANCE the CV to be stronger overall: "
            "improve clarity, sharpen impact statements, highlight transferable skills, "
            "and suggest realistic pivot roles. Preserve the candidate's authentic identity."
        )

    system_msg = (
        "You are a senior career coach and elite resume strategist. "
        "Your job is to improve the candidate's CV while remaining 100% truthful. "
        "Do NOT fabricate experience, skills, or metrics. "
        "Do NOT change the candidate's career identity. "
        "Plain text only — no **, no *, no #, no backticks. ALL CAPS for section headers only."
    )

    user_msg = (
        f"RESUME:\n{resume_text[:2000]}\n\n"
        f"JOB DESCRIPTION:\n{job_text[:1200]}\n\n"
        f"MATCH SCORE: {final_score}%\n"
        f"MISSING SKILLS: {ms}\n\n"
        f"COACHING STRATEGY: {strategy_instruction}\n\n"
        "OUTPUT ALL 5 SECTIONS in this exact format (plain text, no markdown):\n\n"
        "PROFESSIONAL SUMMARY\n"
        "[3-4 powerful sentences.]\n\n"
        "TOP 5 ACHIEVEMENT BULLETS\n"
        "- [Accomplished X as measured by Y by doing Z]\n"
        "- [Leadership or scale impact]\n"
        "- [Tool or technical application]\n"
        "- [Process improvement or collaboration]\n"
        "- [Business growth or customer outcome]\n\n"
        "SKILLS SECTION\n"
        "[Top 15 most relevant skills, comma-separated]\n\n"
        "KEYWORDS TO ADD\n"
        "[8 high-value keywords — only if truthfully applicable]\n\n"
        "CAREER RECOMMENDATION\n"
        "[Honest coaching: best-fit roles, one skill/cert to pursue, pivot advice if low match]\n\n"
        "RULES: No 'I' or 'My'. Plain text only. All 5 sections required."
    )

    try:
        r = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": system_msg},
                {"role": "user",   "content": user_msg}
            ],
            max_tokens=2200,
            temperature=0.4
        )
        return r.choices[0].message.content, None, strategy
    except Exception as e:
        return "", str(e), strategy

# ── RULE-BASED FALLBACK ────────────────────────────────────
def rule_based_aspect_scores(resume_text, job_text, final_score, skill_score, cov_pct):
    scores = {}
    rt = resume_text.lower()
    has_email    = "@" in rt
    has_phone    = bool(re.search(r'\d{3}[-.\s]\d{3,4}', rt))
    has_linkedin = "linkedin" in rt
    has_summary  = any(w in rt for w in ["summary","objective","profile","about"])
    has_exp      = any(w in rt for w in ["experience","work","employment","position"])
    has_edu      = any(w in rt for w in ["education","university","degree","bachelor","master"])
    has_achieve  = any(w in rt for w in ["achieved","increased","reduced","improved","led","managed"])
    has_numbers  = bool(re.search(r'\d+%|\d+ years|\$\d+', rt))
    has_org      = any(w in rt for w in ["organization","club","volunteer","committee","association"])
    has_extra    = any(w in rt for w in ["portfolio","github","project","certification","language"])
    word_count   = len(resume_text.split())
    scores["Overall Impression"]        = min(100, int(final_score * 1.1))
    scores["Contact Information"]       = 60 + (10 if has_email else 0) + (10 if has_phone else 0) + (10 if has_linkedin else 0)
    scores["Relevant Skills"]           = int(skill_score)
    scores["Professional Summary"]      = 70 if has_summary else 35
    scores["Work Experience"]           = 75 if has_exp else 30
    scores["Achievements"]              = (70 if has_achieve else 35) + (10 if has_numbers else 0)
    scores["Education & Certification"] = 75 if has_edu else 40
    scores["Organizational Activity"]   = 70 if has_org else 35
    scores["Consistent & Error-free"]   = 65 if word_count > 200 else 45
    scores["Additional Sections"]       = 70 if has_extra else 40
    scores["Keywords"]                  = int(cov_pct)
    scores["Career Recommendation"]     = int(final_score * 0.9)
    return {k: min(v, 100) for k, v in scores.items()}

def rule_based_rewrite(resume_text, job_text, missing_skills, matched_skills, final_score):
    strategy_note = (
        "✅ Strategy: Tailoring CV to job (match >= 60%)"
        if final_score >= 60 else
        "⚠️ Strategy: Enhancing CV quality (low match — not force-fitting to this job)"
    )
    out = [strategy_note, ""]
    out.append("### PROFESSIONAL SUMMARY")
    skills_str = ', '.join(matched_skills[:4]) if matched_skills else 'relevant fields'
    out.append(
        f"Results-driven professional with proven expertise in {skills_str}. "
        "Demonstrated ability to deliver measurable outcomes through strategic thinking and cross-functional collaboration. "
        "Committed to continuous improvement and driving business value."
    )
    out.append("")
    out.append("### TOP 5 ACHIEVEMENT BULLETS")
    for b in [
        "Accomplished [specific task] as measured by [metric/result] by doing [method/tool used].",
        "Led [project/initiative] resulting in [X]% improvement in [area] through [approach].",
        "Developed [solution/system] that reduced [problem] by [metric] using [technology].",
        "Collaborated with [team/stakeholders] to deliver [outcome] within [timeframe].",
        "Optimized [process/workflow] achieving [result] by implementing [tool/method].",
    ]:
        out.append(f"- {b}")
    out.append("")
    out.append("### SKILLS SECTION")
    all_skills = matched_skills + missing_skills[:5]
    out.append(', '.join(all_skills) if all_skills else "Add relevant technical and soft skills here.")
    out.append("")
    out.append("### KEYWORDS TO ADD")
    out.append(', '.join(missing_skills[:8]) if missing_skills else "Your CV already covers most keywords.")
    out.append("")
    out.append("### CAREER RECOMMENDATION")
    out.append(
        "Based on your skill profile, consider roles aligned with your strongest competencies. "
        "Focus on building portfolio projects that demonstrate end-to-end impact. "
        "Pursue one relevant certification to strengthen your positioning."
    )
    out.append("")
    out.append("---")
    out.append("⚠️ This is a rule-based suggestion. Enable AI for personalized LLM-powered rewrite.")
    return "\n".join(out)

# ── EXPORT ─────────────────────────────────────────────────
def export_docx(resume_text, rewritten_content, final_score, aspect_scores,
                matched_skills, missing_skills, tfidf_score, vector_score, cov_pct, skill_score):
    doc = Document()
    title = doc.add_heading("ResumeFit AI v5.0 — Optimized CV Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading("ATS Match Score Summary", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = "Metric"; hdr[1].text = "Score"
    for m, v in [
        ("Overall ATS Match Score",       f"{final_score}%"),
        ("TF-IDF Similarity (Statistical)", f"{tfidf_score}%"),
        ("Pinecone Semantic Similarity",   f"{vector_score}%"),
        ("Keyword Coverage",               f"{cov_pct}%"),
        ("Skill Match",                    f"{skill_score:.1f}%"),
        ("Matched Skills",                 str(len(matched_skills))),
        ("Missing Skills",                 str(len(missing_skills))),
    ]:
        row = table.add_row().cells
        row[0].text = m; row[1].text = v
    doc.add_paragraph()
    doc.add_heading("Scoring Formula", level=2)
    doc.add_paragraph("Score = 0.4 × TF-IDF + 0.3 × Pinecone + 0.2 × Keyword + 0.1 × Skill")
    if aspect_scores:
        doc.add_heading("12-Aspect CV Analysis", level=1)
        for asp, data in aspect_scores.items():
            p = doc.add_paragraph()
            p.add_run(f"{asp}: {data['score']}/100").bold = True
            doc.add_paragraph(f"  Assessment: {data.get('assessment','')}")
            doc.add_paragraph(f"  Tip: {data.get('tip','')}")
    doc.add_paragraph()
    doc.add_heading("Matched Skills", level=1)
    doc.add_paragraph(', '.join(matched_skills) if matched_skills else "None detected")
    doc.add_heading("Missing Skills", level=1)
    doc.add_paragraph(', '.join(missing_skills) if missing_skills else "None detected")
    doc.add_paragraph()
    if rewritten_content:
        doc.add_heading("AI-Optimized CV Sections", level=1)
        for line in rewritten_content.split("\n"):
            if line.strip():
                if any(line.strip().upper().startswith(x) for x in ["PROFESSIONAL","TOP 5","SKILLS","KEYWORDS","CAREER"]):
                    doc.add_paragraph(line.strip()).runs[0].bold = True
                else:
                    doc.add_paragraph(line.strip())
    doc.add_paragraph()
    doc.add_paragraph("Generated by ResumeFit AI v5.0 | Bernard Lokasasmita | AI Engineering Bootcamp Batch 11")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def export_txt(resume_text, rewritten_content, final_score, matched_skills, missing_skills,
               aspect_scores, tfidf_score, vector_score, cov_pct, skill_score):
    out = [
        "=" * 60,
        "ResumeFit AI v5.0 — Optimized CV Report",
        "=" * 60,
        f"Overall ATS Match Score : {final_score}%",
        "",
        "--- SCORING BREAKDOWN ---",
        f"TF-IDF (Statistical)     : {tfidf_score}%  [weight: 40%]",
        f"Pinecone (Semantic)      : {vector_score}%  [weight: 30%]",
        f"Keyword Coverage         : {cov_pct}%  [weight: 20%]",
        f"Skill Match              : {skill_score:.1f}%  [weight: 10%]",
        "",
        "Formula: Score = 0.4×TF-IDF + 0.3×Pinecone + 0.2×Keyword + 0.1×Skill",
        "",
        "--- 12-ASPECT ANALYSIS ---",
    ]
    if aspect_scores:
        for asp, data in aspect_scores.items():
            out.append(f"\n{asp}: {data['score']}/100")
            out.append(f"  {data.get('assessment','')}")
            out.append(f"  Tip: {data.get('tip','')}")
    out += [
        "",
        "--- MATCHED SKILLS ---",
        ', '.join(matched_skills) if matched_skills else "None",
        "",
        "--- MISSING SKILLS ---",
        ', '.join(missing_skills) if missing_skills else "None",
        "",
        "--- AI-OPTIMIZED SECTIONS ---",
        rewritten_content if rewritten_content else "No rewrite available.",
        "",
        "Generated by ResumeFit AI v5.0 | Bernard Lokasasmita | AI Engineering Bootcamp Batch 11"
    ]
    return "\n".join(out)

# ── SIDEBAR ────────────────────────────────────────────────
with st.sidebar:
    st.image("https://img.icons8.com/fluency/96/resume.png", width=70)
    st.title("ResumeFit AI")
    st.caption("v5.0 — TF-IDF + Pinecone + LLM")
    st.divider()
    st.markdown("### ⚙️ Settings")
    use_ai = st.toggle("Enable AI Analysis", value=True)
    llm_client = None
    llm_model  = GROQ_DEFAULT_MODEL
    use_openai = False

    if use_ai:
        st.markdown("**LLM Provider**")
        use_openai = st.toggle("Switch to ChatGPT (OpenAI)", value=False)
        if not use_openai:
            groq_key = os.environ.get("GROQ_API_KEY", "").strip()
            if groq_key:
                llm_client = get_groq_client(groq_key)
                st.markdown("<div class='success-box'>🟢 <b>Groq ready</b> — Free &amp; Fast ⚡</div>", unsafe_allow_html=True)
            else:
                groq_key_input = st.text_input("Groq API Key (free)", type="password", placeholder="gsk_...",
                                               help="Get your FREE key at https://console.groq.com")
                st.caption("🔑 Free key: [console.groq.com](https://console.groq.com)")
                if groq_key_input:
                    llm_client = get_groq_client(groq_key_input)
                    st.markdown("<div class='success-box'>🟢 <b>Groq ready</b></div>", unsafe_allow_html=True)
                else:
                    st.markdown(
                        "<div class='info-box'>ℹ️ Add <b>GROQ_API_KEY</b> to Colab Secrets for auto-login.</div>",
                        unsafe_allow_html=True
                    )
        else:
            openai_key = st.text_input("OpenAI API Key", type="password", placeholder="sk-...")
            llm_model  = st.selectbox("OpenAI Model", OPENAI_MODELS, index=0)
            if openai_key:
                llm_client = get_openai_client(openai_key)
            else:
                st.markdown("<div class='info-box'>ℹ️ Enter your OpenAI API key above.</div>", unsafe_allow_html=True)
    else:
        st.info("💡 AI OFF — using rule-based engine.")

    # ── NEW: Pinecone section ───────────────────────────────
    st.divider()
    st.markdown("### 🗄️ Vector Database (Pinecone)")
    pinecone_index = None
    embed_model    = None

    pinecone_key = os.environ.get("PINECONE_API_KEY", "").strip()
    if not pinecone_key:
        pinecone_key = st.text_input(
            "Pinecone API Key",
            type="password",
            placeholder="pcsk_...",
            help="Free key at https://pinecone.io"
        )
        st.caption("🔑 Free key: [pinecone.io](https://pinecone.io)")

    if pinecone_key:
        try:
            with st.spinner("Connecting to Pinecone..."):
                pinecone_index = init_pinecone(pinecone_key)
            st.markdown("<div class='success-box'>🟢 <b>Pinecone connected</b> ✅</div>", unsafe_allow_html=True)
            with st.spinner("Loading embedding model..."):
                embed_model = load_embedding_model()
            st.markdown("<div class='vector-box'>🔵 <b>Embeddings ready</b> — all-MiniLM-L6-v2</div>", unsafe_allow_html=True)
        except Exception as e:
            st.markdown(f"<div class='error-box'>❌ Pinecone error: {e}</div>", unsafe_allow_html=True)
    else:
        st.markdown(
            "<div class='info-box'>ℹ️ Add <b>PINECONE_API_KEY</b> to Colab Secrets, "
            "or paste it above. Semantic scoring will be skipped without it.</div>",
            unsafe_allow_html=True
        )

    st.divider()
    st.markdown("### 📊 Hybrid Scoring Formula")
    st.markdown("""

```
Score =
  0.4 × TF-IDF
  0.3 × Pinecone
  0.2 × Keywords
  0.1 × Skills
```

""")
    st.divider()
    st.markdown("### 📋 12 CV Aspects")
    for name, icon, _ in ASPECTS:
        st.markdown(f"{icon} {name}")
    st.divider()
    st.caption("ResumeFit AI v5.0 | Bernard Lokasasmita | Batch 11")

# ── HEADER ─────────────────────────────────────────────────
st.markdown("""
<div class='header-box'>
    <h1 style='color:white;margin:0;font-size:2rem;'>🚀 ResumeFit AI <span style='font-size:1rem;opacity:0.8;'>v5.0</span></h1>
    <p style='color:#e9d5ff;margin:6px 0 0 0;font-size:1rem;'>
        Hybrid AI: TF-IDF + Pinecone Vector DB + 12-Aspect LLM + Smart Rewrite
    </p>
</div>
""", unsafe_allow_html=True)

# Status banners
if use_ai and llm_client:
    provider = "OpenAI" if use_openai else "Groq (Free ⚡)"
    st.markdown(f"<div class='success-box'>🤖 <b>AI ON</b> — {provider} | {llm_model}</div>", unsafe_allow_html=True)
elif use_ai and not llm_client:
    st.markdown("<div class='info-box'>⚠️ <b>AI ON</b> — Enter Groq key in sidebar to activate LLM.</div>", unsafe_allow_html=True)
else:
    st.markdown("<div class='rule-box'>🔧 Rule-Based Mode — Toggle AI in sidebar.</div>", unsafe_allow_html=True)

if pinecone_index:
    st.markdown("<div class='vector-box'>🗄️ <b>Pinecone ACTIVE</b> — Semantic vector similarity enabled.</div>", unsafe_allow_html=True)
else:
    st.markdown("<div class='warn-box'>⚠️ <b>Pinecone NOT connected</b> — Semantic scoring disabled. Add PINECONE_API_KEY to use full hybrid scoring.</div>", unsafe_allow_html=True)

st.markdown("<br>", unsafe_allow_html=True)

# ── INPUT ──────────────────────────────────────────────────
st.markdown("<div class='section-header'>📄 Step 1: Upload CV &amp; Job Description</div>", unsafe_allow_html=True)
col1, col2 = st.columns(2)
with col1:
    st.markdown("**Upload CV / Resume (PDF)**")
    uploaded_file = st.file_uploader("", type=["pdf"], label_visibility="collapsed")
    if uploaded_file:
        st.success(f"✅ {uploaded_file.name}")
with col2:
    st.markdown("**Paste Job Description**")
    job_description = st.text_area("", height=200,
                                   placeholder="Paste the full job description here...",
                                   label_visibility="collapsed")
    if job_description:
        st.caption(f"📝 {len(job_description.split())} words")

st.markdown("<br>", unsafe_allow_html=True)
analyze_btn = st.button("🔍 Analyze My CV — Hybrid AI", use_container_width=True)

# ── ANALYSIS ───────────────────────────────────────────────
if analyze_btn:
    if not uploaded_file:
        st.error("⚠️ Please upload your CV PDF.")
        st.stop()
    if not job_description or len(job_description.strip()) < 50:
        st.error("⚠️ Please paste a job description (min 50 chars).")
        st.stop()
    if use_ai and not llm_client:
        st.error("⚠️ AI enabled but no valid API key. Enter key in sidebar.")
        st.stop()

    with st.spinner("🤖 Running hybrid analysis..."):
        resume_text = extract_text_from_pdf(uploaded_file)
        if not resume_text:
            st.error("Could not extract text. Use a text-based PDF.")
            st.stop()

        # Layer 1 & 2: Statistical scoring
        tfidf_score = calculate_tfidf(resume_text, job_description)
        cov_pct, matched_kw, missing_kw = keyword_coverage(resume_text, job_description)
        rs, js, matched_s, missing_s, extra_s = skill_analysis(resume_text, job_description)
        skill_score = (len(matched_s) / max(len(js), 1)) * 100

        # ── NEW Layer 3: Pinecone — Ingestion + Embedding + Query
        vector_score = 0.0
        if pinecone_index and embed_model:
            with st.spinner("📥 Step 1/2 — Program Ingestion: embedding Job Description → storing in Pinecone Vector DB..."):
                ingest_job_description(job_description, pinecone_index, embed_model)
            st.markdown("<div class='success-box'>✅ <b>Ingestion complete</b> — Job Description embedded &amp; stored in Pinecone.</div>", unsafe_allow_html=True)
            with st.spinner("🔍 Step 2/2 — Model Embedding: embedding CV → querying Pinecone for semantic similarity..."):
                vector_score = query_with_cv(resume_text, pinecone_index, embed_model)
            st.markdown(f"<div class='vector-box'>🔵 <b>Semantic similarity score: {vector_score}%</b> — CV embedding queried against stored Job Description.</div>", unsafe_allow_html=True)

        # ── NEW: Updated weighted formula ──────────────────
        if vector_score > 0:
            # Full 4-layer hybrid
            final_score = round(
                (tfidf_score * 0.4) +
                (vector_score * 0.3) +
                (cov_pct * 0.2) +
                (skill_score * 0.1),
                1
            )
        else:
            # Fallback: original 3-layer formula (no Pinecone)
            final_score = round(
                (tfidf_score * 0.5) +
                (cov_pct * 0.3) +
                (skill_score * 0.2),
                1
            )

        label, color = score_label(final_score)

        # ── Feature #2: Formatting checks ──────────────────
        with st.spinner("📋 Checking ATS formatting compatibility..."):
            pdf_issues  = check_ats_formatting(uploaded_file)
            text_issues = check_text_ats(resume_text)
            ats_fmt_score = ats_format_score(pdf_issues, text_issues)

        # ── Feature #1: Impact language ────────────────────
        with st.spinner("⚡ Analyzing impact language & action verbs..."):
            found_weak, found_buzz, has_numbers, llm_impact = analyze_impact_language(
                resume_text,
                client=llm_client if (use_ai and llm_client) else None,
                model=llm_model
            )

        aspect_scores = {}
        rewritten_cv  = ""
        rewrite_strategy = ""
        ai_errors = []

        if use_ai and llm_client:
            provider_label = "OpenAI" if use_openai else "Groq"
            with st.spinner(f"🧠 12-aspect analysis via {provider_label}..."):
                raw_aspect, err1 = analyze_12_aspects(resume_text, job_description, llm_client, llm_model)
                if err1:
                    ai_errors.append(f"12-Aspect: {err1}")
                else:
                    aspect_scores = parse_aspect_scores(raw_aspect)
                    if len(aspect_scores) < 12:
                        rb = rule_based_aspect_scores(resume_text, job_description, final_score, skill_score, cov_pct)
                        for k, v in rb.items():
                            aspect_scores.setdefault(k, {"score": v, "assessment": "Rule-based estimate.", "tip": "Re-run for AI tips."})

            with st.spinner(f"✍️ Rewriting CV via {provider_label}..."):
                rewritten_cv, err2, rewrite_strategy = rewrite_cv_sections(
                    resume_text, job_description, missing_s, final_score, llm_client, llm_model
                )
                if err2:
                    ai_errors.append(f"CV Rewrite: {err2}")

            for err in ai_errors:
                st.markdown(f"<div class='error-box'>❌ AI Error: {err}</div>", unsafe_allow_html=True)
        else:
            rb_raw = rule_based_aspect_scores(resume_text, job_description, final_score, skill_score, cov_pct)
            aspect_scores = {k: {"score": v, "assessment": "Rule-based estimate.", "tip": "Enable AI for specific tips."} for k, v in rb_raw.items()}
            rewritten_cv = rule_based_rewrite(resume_text, job_description, missing_s, matched_s, final_score)
            rewrite_strategy = "HIGH_MATCH" if final_score >= 60 else "LOW_MATCH"

    st.success("✅ Analysis Complete!")
    st.divider()

    tab1, tab2, tab3, tab4, tab5 = st.tabs([
        "📊 Dashboard",
        "🔍 12-Aspect Analysis",
        "🤖 AI-Optimized CV",
        "⚡ Impact & Format",
        "📥 Export"
    ])

    # ── TAB 1: DASHBOARD ───────────────────────────────────
    with tab1:
        st.markdown("<div class='section-header'>📊 Your Results Dashboard</div>", unsafe_allow_html=True)
        r1, r2, r3 = st.columns([1.2, 1.2, 1])
        with r1:
            st.plotly_chart(gauge_chart(final_score), use_container_width=True)
        with r2:
            rc = radar_chart(matched_s, missing_s)
            if rc: st.plotly_chart(rc, use_container_width=True)
            else:  st.info("Not enough skills for radar chart.")
        with r3:
            st.markdown("<br>", unsafe_allow_html=True)
            for val, sub in [
                (f"{final_score}%", "Overall Score"),
                (f"{vector_score}%", "Semantic (Pinecone)"),
                (f"{tfidf_score}%", "Statistical (TF-IDF)"),
                (f"{len(matched_s)}/{len(js)}", "Skills Matched"),
            ]:
                st.markdown(f"""
                <div class='metric-card'>
                    <h1>{val}</h1>
                    <p>{sub}</p>
                </div><br>
                """, unsafe_allow_html=True)

        st.divider()

        # ── NEW: Scoring breakdown chart ───────────────────
        st.plotly_chart(
            scoring_breakdown_chart(tfidf_score, vector_score, cov_pct, skill_score),
            use_container_width=True
        )

        # ── NEW: Score comparison insight ──────────────────
        if vector_score > 0:
            diff = round(vector_score - tfidf_score, 1)
            if diff > 10:
                st.markdown(
                    f"<div class='vector-box'>💡 <b>Semantic vs Statistical:</b> Pinecone score ({vector_score}%) is "
                    f"{diff}pts higher than TF-IDF ({tfidf_score}%). Your CV uses different wording but "
                    "similar concepts to the job description — good semantic alignment.</div>",
                    unsafe_allow_html=True
                )
            elif diff < -10:
                st.markdown(
                    f"<div class='warn-box'>💡 <b>Semantic vs Statistical:</b> TF-IDF ({tfidf_score}%) is higher than "
                    f"Pinecone ({vector_score}%). Your CV shares keywords but the overall meaning diverges — "
                    "consider rewriting for better conceptual alignment.</div>",
                    unsafe_allow_html=True
                )
            else:
                st.markdown(
                    f"<div class='success-box'>💡 <b>Semantic vs Statistical:</b> TF-IDF ({tfidf_score}%) and "
                    f"Pinecone ({vector_score}%) are consistent — your CV aligns well both in keywords and meaning.</div>",
                    unsafe_allow_html=True
                )

        st.divider()

        if final_score >= 60:
            st.markdown(f"<div class='success-box'>✅ <b>Strong Match ({final_score}%)</b> — AI will tailor your CV to this specific job.</div>", unsafe_allow_html=True)
        else:
            st.markdown(f"<div class='warn-box'>⚠️ <b>Low Match ({final_score}%)</b> — AI will enhance your CV quality overall and suggest better-aligned roles.</div>", unsafe_allow_html=True)

        col_skills_a, col_skills_b = st.columns(2)
        with col_skills_a:
            st.markdown("**✅ Matched Skills**")
            if matched_s:
                st.markdown(" ".join([f"<span class='skill-found'>{s}</span>" for s in matched_s]), unsafe_allow_html=True)
            else:
                st.warning("No matching skills found.")

        with col_skills_b:
            st.markdown("**❌ Missing Skills — Priority to Add**")
            if missing_s:
                # Rank missing skills: technical > AI/ML > soft
                ai_keywords  = {"llm","rag","vector database","embeddings","langchain","fine-tuning",
                                 "generative ai","prompt engineering","mlops","hugging face","pytorch",
                                 "tensorflow","machine learning","deep learning","openai","groq"}
                tech_keywords = {"python","sql","docker","kubernetes","aws","gcp","azure","git",
                                  "fastapi","spark","airflow","dbt","postgresql","mongodb"}
                def skill_priority(s):
                    if s.lower() in ai_keywords:   return 0
                    if s.lower() in tech_keywords:  return 1
                    return 2
                ranked_missing = sorted(missing_s, key=skill_priority)
                for i, s in enumerate(ranked_missing[:10]):
                    priority = ["🔴 High","🟠 Medium","🟡 Low"][skill_priority(s)]
                    st.markdown(
                        f"<span class='skill-missing'>{i+1}. {s}</span> "
                        f"<small style='color:#6b7280'>{priority}</small>",
                        unsafe_allow_html=True
                    )
                if len(missing_s) > 10:
                    st.caption(f"+ {len(missing_s)-10} more missing skills")
            else:
                st.success("No critical skill gaps!")

        st.divider()
        m1, m2, m3, m4, m5, m6, m7 = st.columns(7)
        m1.metric("Overall Score",   f"{final_score}%")
        m2.metric("TF-IDF",         f"{tfidf_score}%")
        m3.metric("Pinecone",       f"{vector_score}%")
        m4.metric("Keywords",       f"{cov_pct}%")
        m5.metric("Skill Match",    f"{len(matched_s)}/{len(js)}")
        m6.metric("ATS Format",     f"{ats_fmt_score}%")
        m7.metric("Impact Score",   f"{llm_impact.get('score', '—')}{'%' if llm_impact.get('score') else ''}")

    # ── TAB 2: 12-ASPECT ───────────────────────────────────
    with tab2:
        st.markdown("<div class='section-header'>🔍 Bedah 12 Aspek CV Kamu</div>", unsafe_allow_html=True)
        if aspect_scores:
            scores_dict = {k: v["score"] for k, v in aspect_scores.items()}
            st.plotly_chart(aspect_bar_chart(scores_dict), use_container_width=True)
            st.divider()
            cols = st.columns(2)
            for i, (asp_name, data) in enumerate(aspect_scores.items()):
                icon = ASPECTS[i][1] if i < len(ASPECTS) else "📌"
                with cols[i % 2]:
                    sc = data["score"]
                    st.markdown(f"""
                    <div class='aspect-card'>
                        <h4>{icon} {asp_name} — {sc}/100 &nbsp; {badge(sc)}</h4>
                        <p>📋 {data.get('assessment','')}</p>
                        <p>💡 <b>Tip:</b> {data.get('tip','')}</p>
                    </div>
                    """, unsafe_allow_html=True)
        else:
            st.warning("No aspect scores available.")

    # ── TAB 3: AI REWRITE ──────────────────────────────────
    with tab3:
        st.markdown("<div class='section-header'>🤖 AI-Optimized CV Sections</div>", unsafe_allow_html=True)
        if rewritten_cv:
            if rewrite_strategy == "HIGH_MATCH":
                st.markdown(f"<div class='success-box'>✅ <b>Tailored Mode</b> — Match score {final_score}% is strong.</div>", unsafe_allow_html=True)
            else:
                st.markdown(f"<div class='warn-box'>⚠️ <b>Enhancement Mode</b> — Match score {final_score}% is low. CV improved overall, not force-fitted.</div>", unsafe_allow_html=True)
            st.divider()
            bc, ac = st.columns(2)
            with bc:
                st.markdown("**🔴 Original CV (excerpt)**")
                st.markdown(f"<div class='before-box'>{resume_text[:600]}...</div>", unsafe_allow_html=True)
            with ac:
                st.markdown("**🟢 Optimized Sections (preview)**")
                clean_preview = clean_ai_text(rewritten_cv[:600])
                st.markdown(f"<div class='after-box'>{clean_preview}...</div>", unsafe_allow_html=True)
            st.divider()
            st.markdown("### 📋 Full Optimized Sections")
            SECTION_HEADERS = ["PROFESSIONAL SUMMARY","TOP 5 ACHIEVEMENT BULLETS","SKILLS SECTION","KEYWORDS TO ADD","CAREER RECOMMENDATION"]
            SECTION_ICONS = {"PROFESSIONAL SUMMARY":"📝","TOP 5 ACHIEVEMENT BULLETS":"🏆","SKILLS SECTION":"🛠️","KEYWORDS TO ADD":"🔑","CAREER RECOMMENDATION":"🎯"}
            clean_full = clean_ai_text(rewritten_cv)
            for line in clean_full.split("\n"):
                stripped = line.strip()
                if not stripped:
                    st.write("")
                elif stripped.upper() in SECTION_HEADERS:
                    icon = SECTION_ICONS.get(stripped.upper(), "🔹")
                    st.markdown(f"#### {icon} {stripped.title()}")
                    st.markdown("---")
                elif stripped.startswith("- ") or stripped.startswith("• "):
                    st.markdown(f"&nbsp;&nbsp;{stripped}")
                else:
                    st.write(stripped)
        else:
            st.warning("No rewrite available.")

    # ── TAB 4: IMPACT & FORMAT ─────────────────────────────
    with tab4:
        st.markdown("<div class='section-header'>⚡ Impact Language & ATS Formatting</div>", unsafe_allow_html=True)

        col_imp, col_fmt = st.columns(2)

        # ── Impact Language ────────────────────────────────
        with col_imp:
            st.markdown("### ⚡ Impact Language Analysis")

            # Score badge
            impact_score = llm_impact.get("score")
            if impact_score is not None:
                clr = "#22c55e" if impact_score >= 75 else "#f59e0b" if impact_score >= 50 else "#ef4444"
                st.markdown(
                    f"<div class='metric-card'><h1 style='color:{clr}'>{impact_score}/100</h1>"
                    f"<p>Impact Language Score</p></div><br>",
                    unsafe_allow_html=True
                )

            # Quantification check
            is_quantified = llm_impact.get("quantified", has_numbers)
            if is_quantified:
                st.markdown("<div class='success-box'>✅ <b>Quantified achievements detected</b> — numbers/metrics found in CV.</div>", unsafe_allow_html=True)
            else:
                st.markdown("<div class='error-box'>❌ <b>No quantified achievements</b> — add metrics (%, $, numbers) to bullet points.</div>", unsafe_allow_html=True)

            st.markdown("<br>", unsafe_allow_html=True)

            # Weak verbs
            st.markdown("**🔴 Weak Phrases to Replace**")
            if found_weak:
                for v in found_weak:
                    st.markdown(f"<span class='skill-missing'>❌ \"{v}\"</span>", unsafe_allow_html=True)
                st.caption("Replace with strong action verbs: Led, Built, Reduced, Delivered, Achieved, Grew, Launched...")
            else:
                st.markdown("<div class='success-box'>✅ No weak phrases detected.</div>", unsafe_allow_html=True)

            st.markdown("<br>**🟡 Buzzwords to Remove**")
            if found_buzz:
                for b in found_buzz:
                    st.markdown(f"<span class='score-badge-yellow'> ⚠️ \"{b}\"</span>", unsafe_allow_html=True)
                st.caption("These words are meaningless to recruiters and ATS. Remove or replace with concrete evidence.")
            else:
                st.markdown("<div class='success-box'>✅ No buzzwords detected.</div>", unsafe_allow_html=True)

            # LLM tip
            tip = llm_impact.get("tip")
            if tip:
                st.markdown("<br>", unsafe_allow_html=True)
                st.markdown(f"<div class='strategy-box'>💡 <b>Top Tip:</b> {tip}</div>", unsafe_allow_html=True)

        # ── ATS Formatting ─────────────────────────────────
        with col_fmt:
            st.markdown("### 📋 ATS Formatting Check")

            # Format score
            fmt_clr = "#22c55e" if ats_fmt_score >= 80 else "#f59e0b" if ats_fmt_score >= 55 else "#ef4444"
            st.markdown(
                f"<div class='metric-card'><h1 style='color:{fmt_clr}'>{ats_fmt_score}/100</h1>"
                f"<p>ATS Compatibility Score</p></div><br>",
                unsafe_allow_html=True
            )

            # PDF structure issues
            st.markdown("**📄 PDF Structure**")
            for label_txt, detail, sev in pdf_issues:
                if sev == "error":
                    st.markdown(f"<div class='error-box'>{label_txt}<br><small>{detail}</small></div>", unsafe_allow_html=True)
                elif sev == "warning":
                    st.markdown(f"<div class='warn-box'>{label_txt}<br><small>{detail}</small></div>", unsafe_allow_html=True)
                else:
                    st.markdown(f"<div class='success-box'>{label_txt}</div>", unsafe_allow_html=True)

            st.markdown("<br>**📝 Content Structure**")
            for label_txt, detail, sev in text_issues:
                if sev == "error":
                    st.markdown(f"<div class='error-box'>{label_txt}<br><small>{detail}</small></div>", unsafe_allow_html=True)
                elif sev == "warning":
                    st.markdown(f"<div class='warn-box'>{label_txt}<br><small>{detail}</small></div>", unsafe_allow_html=True)
                else:
                    st.markdown(f"<div class='success-box'>{label_txt}</div>", unsafe_allow_html=True)

    # ── TAB 5: EXPORT ──────────────────────────────────────
    with tab5:
        st.markdown("<div class='section-header'>📥 Export Your Optimized CV Report</div>", unsafe_allow_html=True)
        e1, e2 = st.columns(2)
        with e1:
            st.markdown("### 📄 Export as DOCX")
            docx_buf = export_docx(
                resume_text, rewritten_cv, final_score, aspect_scores,
                matched_s, missing_s, tfidf_score, vector_score, cov_pct, skill_score
            )
            st.download_button(
                label="📥 Download Optimized CV (.docx)",
                data=docx_buf,
                file_name="resumefit_optimized_cv_v5.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
        with e2:
            st.markdown("### 📋 Export as TXT")
            txt_report = export_txt(
                resume_text, rewritten_cv, final_score, matched_s, missing_s,
                aspect_scores, tfidf_score, vector_score, cov_pct, skill_score
            )
            st.download_button(
                label="📥 Download Report (.txt)",
                data=txt_report,
                file_name="resumefit_report_v5.txt",
                mime="text/plain",
                use_container_width=True
            )
        st.divider()
        m1, m2, m3, m4, m5 = st.columns(5)
        m1.metric("Overall Score", f"{final_score}%")
        m2.metric("TF-IDF",       f"{tfidf_score}%")
        m3.metric("Pinecone",     f"{vector_score}%")
        m4.metric("Keywords",     f"{cov_pct}%")
        m5.metric("Aspects",      "12")
